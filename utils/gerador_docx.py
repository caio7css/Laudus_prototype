"""
adicionar validação para poder alterar somente campo específico, ou seja
se o médico mencionar apenas um ponto do laudo, alterar somente ela e manter o laudo com as informações em estado normal.
"""

import io
import os
import unicodedata

from docx import Document


def _normalizar(texto: str) -> str:
    # Remove acentos e caixa, para comparação tolerante a maiúsculas/acentos.
    texto = unicodedata.normalize('NFKD', texto)
    texto = ''.join(c for c in texto if not unicodedata.combining(c))
    return texto.lower().strip()


def _texto_normalizado_com_posicoes(texto: str) -> tuple[str, list[int]]:
    """Normaliza o texto e conserva o índice original de cada caractere."""
    normalizado = []
    posicoes_originais = []
    for indice, caractere in enumerate(texto):
        parte = unicodedata.normalize('NFKD', caractere)
        parte = ''.join(c for c in parte if not unicodedata.combining(c)).lower()
        normalizado.append(parte)
        posicoes_originais.extend([indice] * len(parte))
    return ''.join(normalizado), posicoes_originais


def extrair_secoes(texto_transcrito: str, placeholders: list[str]) -> dict:
    """
    Divide o texto ditado em seções, uma para cada placeholder.

    O médico dita o nome do placeholder em voz alta (ex.: "Transparência
    pulmonar") seguido do achado. Esta função localiza cada menção a um
    placeholder no texto e atribui a ele todo o conteúdo até a próxima
    menção encontrada.

    Placeholders não mencionados no texto retornam string vazia (o parágrafo
    correspondente do template é então mantido como está).
    """
    texto_norm, posicoes_originais = _texto_normalizado_com_posicoes(texto_transcrito)

    posicoes = []
    for ph in placeholders:
        ph_norm = _normalizar(ph)
        idx = texto_norm.find(ph_norm)
        if idx != -1:
            inicio = posicoes_originais[idx]
            fim = posicoes_originais[idx + len(ph_norm) - 1] + 1
            posicoes.append((inicio, fim, ph))

    posicoes.sort(key=lambda x: x[0])

    secoes = {ph: "" for ph in placeholders}
    for i, (_inicio, fim_placeholder, ph) in enumerate(posicoes):
        inicio_conteudo = fim_placeholder
        fim = posicoes[i + 1][0] if i + 1 < len(posicoes) else len(texto_transcrito)
        trecho = texto_transcrito[inicio_conteudo:fim].strip(" :.,-\n\t")
        if trecho:
            secoes[ph] = trecho

    return secoes


def _definir_texto_paragrafo(paragrafo, novo_texto: str):
    """Substitui o texto de um parágrafo mantendo a formatação do 1º run."""
    if not paragrafo.runs:
        paragrafo.add_run(novo_texto)
        return
    paragrafo.runs[0].text = novo_texto
    for run in paragrafo.runs[1:]:
        run.text = ""


def _capitalizar(frase: str) -> str:
    frase = frase.strip()
    if not frase:
        return frase
    if frase[-1] not in ".!?":
        frase += "."
    return frase[0].upper() + frase[1:]


def gerar_laudo(
    caminho_template: str,
    dados_paciente: dict,
    secoes: dict,
    placeholders: list[str],
    medico: dict | None = None,
) -> io.BytesIO:
    """
    Preenche o template .docx e devolve o arquivo pronto em memória (BytesIO).

    dados_paciente: {"nome", "prontuario", "nascimento", "data_realizacao"}
    secoes: dicionário {placeholder: achado_ditado}, geralmente vindo de extrair_secoes()
    placeholders: lista ordenada dos placeholders válidos para este laudo
    medico: {"nome", "crm", "uf"} (uf é opcional, default "PB")
    """
    doc = Document(caminho_template)

    for paragrafo in doc.paragraphs:
        texto = paragrafo.text.strip()
        texto_norm = _normalizar(texto)

        if texto_norm.startswith('nome:'):
            novo = f"NOME: {dados_paciente.get('nome', '')}     REG: {dados_paciente.get('prontuario', '')}"
            _definir_texto_paragrafo(paragrafo, novo)

        elif texto_norm.startswith('data de nascimento:'):
            _definir_texto_paragrafo(
                paragrafo, f"DATA DE NASCIMENTO: {dados_paciente.get('nascimento', '')}"
            )

        elif texto_norm.startswith('data de realizacao:'):
            _definir_texto_paragrafo(
                paragrafo, f"DATA DE REALIZAÇÃO: {dados_paciente.get('data_realizacao', '')}"
            )

        elif texto.startswith('-'):
            conteudo_norm = _normalizar(texto.lstrip('- ').strip())
            for ph in placeholders:
                if conteudo_norm.startswith(_normalizar(ph)):
                    achado = secoes.get(ph, '').strip()
                    if achado:
                        _definir_texto_paragrafo(
                            paragrafo, f"- {ph}: {_capitalizar(achado)}"
                        )
                    break

        elif medico and texto.upper().startswith('DR'):
            _definir_texto_paragrafo(paragrafo, medico.get('nome', 'DR(A).'))

        elif medico and texto.upper().startswith('CRM'):
            uf = medico.get('uf', 'PB')
            _definir_texto_paragrafo(paragrafo, f"CRM-{uf} {medico.get('crm', '')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def caminho_preview_pdf(caminho_docx: str) -> str:
    """
    Convenção do projeto: para cada 'NOME_DO_LAUDO.docx' existe um
    'NOME_DO_LAUDO_pdf.pdf' irmão, na mesma pasta, usado apenas para
    pré-visualização (o PyMuPDF não abre .docx diretamente).
    """
    base, _ext = os.path.splitext(caminho_docx)
    return f"{base}_pdf.pdf"
